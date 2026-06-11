"""
Word-PDF 相互转换工具
支持：.docx → PDF、.pdf → .docx
"""
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os, sys, threading

class ConverterApp:
    def __init__(self):
        self.win = tk.Tk()
        self.win.title('Word ↔ PDF 转换工具')
        self.win.geometry('520x380')
        self.win.resizable(False, False)

        # Title
        tk.Label(self.win, text='Word ↔ PDF 相互转换', font=('微软雅黑', 16, 'bold')).pack(pady=15)
        tk.Label(self.win, text='支持 .docx → PDF 和 .pdf → .docx', font=('微软雅黑', 10), fg='gray').pack()

        # File selection
        frm = tk.Frame(self.win)
        frm.pack(pady=20)
        tk.Label(frm, text='选择文件：', font=('微软雅黑', 11)).pack(side=tk.LEFT)
        self.file_path = tk.StringVar()
        tk.Entry(frm, textvariable=self.file_path, width=40, font=('微软雅黑', 10)).pack(side=tk.LEFT, padx=5)
        tk.Button(frm, text='浏览', command=self.pick_file, width=8, font=('微软雅黑', 10)).pack(side=tk.LEFT)

        # Convert buttons
        btn_frm = tk.Frame(self.win)
        btn_frm.pack(pady=15)

        self.btn_w2p = tk.Button(btn_frm, text='Word → PDF', command=lambda: self.convert('w2p'),
                                  width=15, height=2, font=('微软雅黑', 12, 'bold'), bg='#2196F3', fg='white')
        self.btn_w2p.pack(side=tk.LEFT, padx=10)

        self.btn_p2w = tk.Button(btn_frm, text='PDF → Word', command=lambda: self.convert('p2w'),
                                  width=15, height=2, font=('微软雅黑', 12, 'bold'), bg='#4CAF50', fg='white')
        self.btn_p2w.pack(side=tk.LEFT, padx=10)

        # Progress
        self.progress = ttk.Progressbar(self.win, mode='indeterminate', length=400)
        self.progress.pack(pady=10)
        self.status = tk.Label(self.win, text='就绪', font=('微软雅黑', 10), fg='gray')
        self.status.pack()

        self._result = ''

        # Output info
        self.info = tk.Label(self.win, text='Word→PDF 需要安装 Microsoft Word\nPDF→Word 无需额外依赖',
                              font=('微软雅黑', 9), fg='#888')
        self.info.pack(pady=15)

    def pick_file(self):
        path = filedialog.askopenfilename(
            filetypes=[('Word/PDF 文件', '*.docx;*.pdf'), ('Word 文件', '*.docx'), ('PDF 文件', '*.pdf')])
        if path:
            self.file_path.set(path)
            ext = os.path.splitext(path)[1].lower()
            if ext == '.docx':
                self.status.config(text='已选 Word 文件，点击 Word → PDF 转换')
            else:
                self.status.config(text='已选 PDF 文件，点击 PDF → Word 转换')

    def convert(self, direction):
        path = self.file_path.get()
        if not path:
            messagebox.showwarning('提示', '请先选择文件')
            return
        self.progress.start()
        self.status.config(text='转换中...')
        self.btn_w2p.config(state='disabled')
        self.btn_p2w.config(state='disabled')

        def run():
            try:
                if direction == 'w2p':
                    self._word_to_pdf(path)
                else:
                    self._pdf_to_word(path)
            except Exception as e:
                err_msg = str(e)
                self.win.after(0, lambda msg=err_msg: messagebox.showerror('错误', msg))
            finally:
                self.win.after(0, self._done)

        threading.Thread(target=run, daemon=True).start()

    def _word_to_pdf(self, path):
        from docx2pdf import convert
        out_dir = os.path.dirname(path)
        # 修复 PyInstaller noconsole 模式下 sys.stderr 为 None 导致 tqdm 报错
        import io
        _stderr_backup = sys.stderr
        sys.stderr = io.StringIO()
        try:
            convert(path, out_dir)
        finally:
            sys.stderr = _stderr_backup
        pdf_path = os.path.splitext(path)[0] + '.pdf'
        self._result = f'转换成功！\n{os.path.basename(pdf_path)}'

    def _pdf_to_word(self, path):
        import fitz
        fitz.Rect.get_area = lambda self: abs(self)  # 兼容新版 PyMuPDF
        from pdf2docx import Converter
        docx_path = os.path.splitext(path)[0] + '.docx'
        cv = Converter(path)
        cv.convert(docx_path)
        cv.close()
        self._normalize_docx_fonts(docx_path)
        self._result = f'转换成功！\n{os.path.basename(docx_path)}'

    def _normalize_docx_fonts(self, docx_path):
        """修正 PDF 转换后的乱码字体名，映射为 Windows 标准字体"""
        from docx import Document
        from docx.oxml.ns import qn

        # 字体映射规则：将 Adobe 嵌入字体名 → Windows 标准字体
        FONT_MAP = {
            'AdvTimes': 'Times New Roman',
            'AdvMathPi1': 'Symbol',
        }

        def normalize(name):
            if not name:
                return 'Times New Roman'
            if name in FONT_MAP:
                return FONT_MAP[name]
            # AdvP*, AdvPS* 等 Adobe Pi/PiStd 字体 → Times New Roman
            if name.startswith('AdvP') or name.startswith('AdvM'):
                return 'Times New Roman'
            return name

        def fix_paragraph(p):
            for run in p.runs:
                old_name = run.font.name
                new_name = normalize(old_name)
                if old_name != new_name:
                    run.font.name = new_name
                    # 同步修正 east-asia 字体属性，否则 CJK 字符仍用旧字体
                    rpr = run._element.rPr
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            rfonts.set(qn('w:eastAsia'), new_name)

        doc = Document(docx_path)

        # 正文段落
        for p in doc.paragraphs:
            fix_paragraph(p)

        # 表格中的文字
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        fix_paragraph(p)

        doc.save(docx_path)

    def _done(self):
        self.progress.stop()
        self.btn_w2p.config(state='normal')
        self.btn_p2w.config(state='normal')
        self.status.config(text=self._result)

    def run(self):
        self.win.mainloop()

if __name__ == '__main__':
    ConverterApp().run()
